"""
Create the final GSIP research paper in Word format.

This version creates a proper academic paper with:
- Clear methodology explanation
- Code examples
- Proper structure
- Real technical depth

Requirements:
    pip install python-docx

Usage:
    python scripts/create_final_paper.py
"""

import os
import sys
import re

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def create_final_paper():
    """Create the comprehensive GSIP Research Paper."""
    
    # Read the markdown file
    md_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'GSIP_Research_Paper_Final.md')
    
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    doc = Document()
    
    # Set document styles
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15
    style.paragraph_format.space_after = Pt(8)
    
    # Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Times New Roman'
        heading_style.font.bold = True
        if i == 1:
            heading_style.font.size = Pt(16)
        elif i == 2:
            heading_style.font.size = Pt(14)
        else:
            heading_style.font.size = Pt(12)
    
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # Parse markdown and convert to docx
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    current_table_rows = []
    in_table = False
    
    while i < len(lines):
        line = lines[i]
        
        # Handle code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block - add all code lines
                for code_line in code_lines:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    run = p.add_run(code_line)
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
                doc.add_paragraph()  # Space after code
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                current_table_rows = []
            
            # Parse table row
            cells = [c.strip() for c in line.split('|')[1:-1]]
            
            # Skip separator rows (containing only dashes)
            if cells and not all(set(c) <= set('-: ') for c in cells):
                current_table_rows.append(cells)
            
            i += 1
            continue
        elif in_table:
            # End of table
            if current_table_rows:
                n_cols = max(len(row) for row in current_table_rows)
                n_rows = len(current_table_rows)
                
                if n_rows > 0 and n_cols > 0:
                    table = doc.add_table(rows=n_rows, cols=n_cols)
                    table.style = 'Table Grid'
                    
                    for row_idx, row_data in enumerate(current_table_rows):
                        for col_idx, cell_data in enumerate(row_data):
                            if col_idx < n_cols:
                                table.rows[row_idx].cells[col_idx].text = cell_data
                                if row_idx == 0:
                                    # Header row
                                    table.rows[row_idx].cells[col_idx].paragraphs[0].runs[0].bold = True
                                    set_cell_shading(table.rows[row_idx].cells[col_idx], 'E0E0E0')
                    
                    doc.add_paragraph()  # Space after table
            
            current_table_rows = []
            in_table = False
            # Don't increment i, process this line normally
        
        # Skip empty lines in some contexts
        if not line.strip():
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# ') and not line.startswith('## '):
            # Main title or H1
            text = line[2:].strip()
            if 'GSIP' in text and 'General' in text:
                # Title
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.size = Pt(18)
            else:
                doc.add_heading(text, level=1)
            i += 1
            continue
        
        if line.startswith('## '):
            text = line[3:].strip()
            doc.add_heading(text, level=2)
            i += 1
            continue
        
        if line.startswith('### '):
            text = line[4:].strip()
            doc.add_heading(text, level=3)
            i += 1
            continue
        
        # Handle horizontal rules
        if line.strip() == '---':
            doc.add_paragraph('─' * 50)
            i += 1
            continue
        
        # Handle bullet points
        if line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            p = doc.add_paragraph(style='List Bullet')
            
            # Handle bold text
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            
            i += 1
            continue
        
        # Handle numbered lists
        match = re.match(r'^(\d+)\.\s+(.+)$', line.strip())
        if match:
            text = match.group(2)
            p = doc.add_paragraph(style='List Number')
            
            # Handle bold text
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            
            i += 1
            continue
        
        # Handle blockquotes
        if line.strip().startswith('> '):
            text = line.strip()[2:]
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            i += 1
            continue
        
        # Regular paragraph
        text = line.strip()
        if text:
            p = doc.add_paragraph()
            
            # Handle inline code
            text = re.sub(r'`([^`]+)`', r'[\1]', text)
            
            # Handle bold text
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
        
        i += 1
    
    # Save document
    output_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'GSIP_Research_Paper_Final.docx')
    doc.save(output_path)
    print(f"Research paper saved to: {output_path}")
    
    # Print stats
    word_count = len(content.split())
    print(f"Approximate word count: {word_count}")
    
    return output_path


if __name__ == "__main__":
    create_final_paper()
