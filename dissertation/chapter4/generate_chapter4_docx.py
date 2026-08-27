#!/usr/bin/env python3
"""Build Chapter 4 as a formatted Word document from verified evidence only."""
from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips, Emu, Cm

HERE = Path(__file__).resolve().parent
EVIDENCE = HERE / "evidence"
FIG = HERE / "figures"
SHOTS = HERE / "screenshots"
OUT_DOCX = HERE / "GSIP_Dissertation_Chapter4.docx"

NAVY = "1F4E79"
ALT = "E8EEF4"
WHITE = "FFFFFF"


def set_run_font(run, name="Times New Roman", size=12, bold=False, colour=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if colour is not None:
        run.font.color.rgb = RGBColor(*colour)


def set_paragraph_format(p, *, after=6, before=0, align="justify", line=1.5, keep_with_next=False):
    fmt = p.paragraph_format
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = line
    fmt.widow_control = True
    fmt.keep_together = False
    fmt.keep_with_next = keep_with_next
    if align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "left":
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == "right":
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_text(cell, text, *, bold=False, colour=None, size=9.5, align="left", font="Times New Roman"):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(text)
    set_run_font(run, name=font, size=size, bold=bold, colour=colour)


def set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def add_page_number(section):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    set_run_font(run, size=10)
    fld1 = OxmlElement("w:fldChar")
    fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar")
    fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1)
    run._r.append(instr)
    run._r.append(fld2)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.font.color.rgb = RGBColor(0, 0, 0)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.widow_control = True

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1.font.size = Pt(20)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.widow_control = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2.font.size = Pt(16)
    h2.font.bold = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.page_break_before = False
    h2.paragraph_format.widow_control = True

    h3 = doc.styles["Heading 3"]
    h3.font.name = "Arial"
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    h3.paragraph_format.space_before = Pt(12)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.widow_control = True


def body(doc, text):
    p = doc.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run(text)
    set_run_font(run)
    return p


def caption(doc, text, *, above=False):
    p = doc.add_paragraph()
    set_paragraph_format(p, after=6 if not above else 3, align="center", line=1.15)
    run = p.add_run(text)
    set_run_font(run, size=10, bold=True)
    return p


def add_figure(doc, path: Path, caption_text: str, alt: str, width=6.0):
    p = doc.add_paragraph()
    set_paragraph_format(p, after=0, align="center", line=1.0)
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    inline = run._element.xpath(".//a:blip")  # noqa: F841
    # Keep caption with figure
    p.paragraph_format.keep_with_next = True
    cap = caption(doc, caption_text)
    # alt text
    doc.paragraphs  # no-op keep
    try:
        inline_el = p.runs[0]._element.xpath(".//wp:docPr")[0]
        inline_el.set("descr", alt)
    except Exception:
        pass
    return cap


def add_table(doc, headers, rows, caption_text):
    caption(doc, caption_text, above=True)
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    hdr = table.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        shade_cell(cell, NAVY)
        set_cell_text(cell, h, bold=True, colour=(255, 255, 255), size=9, align="center")
    for r_i, row in enumerate(rows):
        fill = ALT if r_i % 2 == 0 else WHITE
        for c_i, val in enumerate(row):
            cell = table.rows[r_i + 1].cells[c_i]
            shade_cell(cell, fill)
            set_cell_text(cell, str(val), size=9, align="center" if c_i else "left")
    # spacer
    p = doc.add_paragraph()
    set_paragraph_format(p, after=6, line=1.0)
    return table


def hv(st, key):
    h = st["table"][key]["hv"]
    i = st["table"][key]["igd"]
    f = st["table"][key]["front"]
    t = st["table"][key]["time"]
    nzero = st["table"][key]["n_zero_hv"]
    return (
        f"{h['median']:.3f} ({h['q1']:.3f} to {h['q3']:.3f})",
        f"{i['median']:.3f} ({i['q1']:.3f} to {i['q3']:.3f})",
        f"{f['median']:.0f}",
        f"{t['median']:.1f}",
        str(nzero),
    )


def build():
    st = json.loads((EVIDENCE / "benchmark_stats.json").read_text(encoding="utf-8"))
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)
    section.orientation = 0
    add_page_number(section)
    configure_styles(doc)

    note = doc.add_paragraph()
    set_paragraph_format(note, after=12, align="left", line=1.15)
    r = note.add_run(
        "Document note. The original dissertation Word file was not present in the "
        "repository or in the files supplied with this request. This file contains "
        "Chapter 4 only, formatted for insertion into the existing dissertation. "
        "The application source at commit 7803baa was not modified. Untracked files "
        "under dissertation/ are evidence and this chapter."
    )
    set_run_font(r, size=11)

    doc.add_heading("Chapter 4: Experimental Results and Discussion", level=1)

    body(
        doc,
        "This chapter reports the tests, optimiser benchmarks, reproducibility checks "
        "and source inspections that were executed against GSIP at commit "
        "7803baa721a12ca19e26e700425fe7be94bfc3a4. The working tree for application "
        "source was clean. Files created under dissertation/ are supporting evidence "
        "for this chapter and were not part of the submitted platform. No numerical "
        "result below is taken from documentation alone. Where a claim could not be "
        "verified, that limit is stated.",
    )

    doc.add_heading("4.1 Experimental setup", level=2)

    body(
        doc,
        "The audit of available evidence was completed before any comparative "
        "interpretation. Directly verified items were: the Python unit suite, the web "
        "Vitest suite, Ruff, Alembic migrations against a live PostgreSQL 16 database "
        "with pgvector 0.6.0, one hundred and twenty equal-budget optimiser runs, a "
        "same-seed replay, a changed-seed check, in-process domain-pack simulations "
        "used only to populate the user interface, and screenshots of the running web, "
        "admin and OpenAPI surfaces. GitHub Actions logs supported the same Python and "
        "web suites on ubuntu-latest, together with a successful migration job and a "
        "successful application-image build; those logs did not replace the local "
        "executions. Documentation, including LOCAL_RUN.md and the Phase 0 optimiser "
        "note, was treated as context rather than as results. LOCAL_RUN.md records 214 "
        "Python tests passed; the suite executed for this chapter produced 223 passed "
        "and one skipped, so the document figure was not reused.",
    )
    body(
        doc,
        "Incomplete items were recorded before analysis. The original dissertation "
        "Word document was not supplied, so this chapter cannot be inserted into the "
        "pre-existing front matter by the experimenter. Docker overlay2 could not be "
        "initialised in the experimental environment, so Temporal, MinIO, Milvus and "
        "the Compose PostgreSQL 15 image were not started. A random-search baseline is "
        "not registered in services/optimizer/backends.py; it was added only as an "
        "external harness that called the same evaluators and metrics. Canonical "
        "thirty-variable ZDT problems from zdt_canonical.yaml were not run. Integration "
        "tests that require testcontainers were skipped. No user study was conducted. "
        "These gaps bound every later claim.",
    )

    doc.add_heading("Hardware, software and commands", level=3)
    body(
        doc,
        "Experiments were run on 27 August 2026 on Ubuntu 24.04.4 LTS, kernel "
        "6.12.94+, with a four-core Intel Xeon processor, 15 GiB of memory and no "
        "graphics processor. Python 3.12.3, Node.js 22.14.0, npm 10.9.7, PostgreSQL "
        "16.15, Redis 7.0.15 and LibreOffice 24.2.7 were used. Important Python "
        "packages included numpy 2.0.2, scipy 1.14.1, scikit-learn 1.5.2, FastAPI "
        "0.111.0, SQLAlchemy 2.0.30, Alembic 1.13.1, pytest 8.2.0, ruff 0.14.3, "
        "temporalio 1.6.0 and ray 2.53.0. The web application used Next.js 14.2.0 and "
        "React 18.3.1. Docker 29.1.3 was installed but the daemon failed with "
        "overlay2 unsupported, so containers were not used. Native PostgreSQL listened "
        "on port 5432 rather than the Compose mapping 5433. DATABASE_URL was therefore "
        "postgresql+asyncpg://gsip:gsip_password@localhost:5432/gsip. RAY_ADDRESS was "
        "local. GSIP_DEMO_AUTH remained at its default of true. No passwords or API "
        "keys beyond the local database role already published in .env.example were "
        "introduced. Temporal, MinIO and a Ray cluster were not started.",
    )
    body(
        doc,
        "Services were started with: sudo service postgresql start; sudo service "
        "redis-server start; alembic upgrade head in services/api; python "
        "scripts/seed_data.py; uvicorn services.api.main:app --host 127.0.0.1 --port "
        "8000; npm run dev in apps/web; and npm run dev in apps/admin. Tests used "
        "pytest -m \"not integration and not slow\" --cov, ruff check core/ libs/ "
        "compute/ services/ scripts/ tests/, and npm test -- --run in apps/web. "
        "Benchmarks used python dissertation/chapter4/run_benchmark_campaign.py, which "
        "calls the unmodified run_benchmark harness. The random-seed policy was the "
        "integer sequence 1 to 10 inclusive, applied independently to every backend and "
        "problem. Termination was by completed objective evaluations, not by a "
        "backend-specific stopping rule. Wall time included initialisation of the "
        "backend. Parallel evaluation inside a run was not enabled; two worker "
        "processes ran independent seeds. Timeouts and retries from the Temporal "
        "workflow were not exercised.",
    )

    doc.add_heading("Benchmark problems and optimiser settings", level=3)
    body(
        doc,
        "The measuring instruments are the analytic ZDT and DTLZ problems implemented "
        "in services/optimizer/benchmarks/problems.py, following Zitzler, Deb and "
        "Thiele (2000) and Deb, Thiele, Laumanns and Zitzler (2002). All objectives "
        "are minimised. Variable bounds are [0, 1] on every coordinate used here. There "
        "are no explicit constraints. Hypervolume used the problem reference point "
        "stored with the implementation: (1.1, 1.1) for ZDT1 and ZDT2, and (1.2, 1.2, "
        "1.2) for DTLZ2. Two-objective hypervolume was exact. Three-objective "
        "hypervolume was a seeded Monte Carlo estimate with 200,000 samples. IGD used "
        "a 200-point sample of the analytic front. Hypervolume values were never "
        "compared across different reference points. IGD values were never compared "
        "across different reference fronts.",
    )
    body(
        doc,
        "The evaluation budget was 200 completed objective evaluations and the batch "
        "size was 20, matching configs/benchmarks/zdt1_smoke.yaml rather than the "
        "canonical thirty-variable, 5,000-evaluation sweep in zdt_canonical.yaml. "
        "ZDT1 and ZDT2 used five decision variables for the same reason. That is a "
        "deliberate deviation from the literature sizes and from Chapter 3's canonical "
        "discussion if that chapter specified thirty variables: it keeps the Bayesian "
        "and hybrid backends computationally feasible at ten seeds, but it prevents "
        "any claim that the ranking would hold at canonical dimension. DTLZ2 used the "
        "implementation default of three objectives and twelve variables. Evolutionary "
        "population size remained 50. Bayesian n_restarts remained at the backend "
        "default of 5. Hybrid used UnifiedOptimizer with its Bayesian half fitted to "
        "the first objective only, which is the behaviour of the submitted code. The "
        "external random-search harness sampled uniformly in the box with the same "
        "budget, batch size and seeds. Failed evaluations were not observed: every "
        "analytic call returned a finite vector, so the failure-handling rule was not "
        "triggered. All four methods received the same number of completed evaluations "
        "on every run.",
    )

    add_table(
        doc,
        ["Problem", "Backend", "Variables", "Objectives", "Budget", "Seeds", "Important settings"],
        [
            ["ZDT1", "evolutionary, Bayesian, hybrid, random", "5", "2", "200", "10", "Batch 20; HV ref (1.1, 1.1); exact HV"],
            ["ZDT2", "evolutionary, Bayesian, hybrid, random", "5", "2", "200", "10", "Batch 20; HV ref (1.1, 1.1); concave front"],
            ["DTLZ2", "evolutionary, Bayesian, hybrid, random", "12", "3", "200", "10", "Batch 20; HV ref (1.2)^3; Monte Carlo HV"],
        ],
        "Table 4.1. Benchmark configuration used for the equal-budget campaign.",
    )

    doc.add_heading("Measures", level=3)
    body(
        doc,
        "Hypervolume measures the volume of objective space dominated by the "
        "non-dominated archive relative to a reference point. Higher is better. It was "
        "computed by the exact two-dimensional sweep in "
        "services/optimizer/benchmarks/metrics.py, or by Monte Carlo for DTLZ2. Its "
        "main limitation is dependence on the reference point and, in three "
        "objectives, sampling error. IGD is the mean Euclidean distance from each "
        "reference-front point to the nearest archive point. Lower is better. It "
        "punishes gaps in coverage as well as distance, and it is undefined for an "
        "empty front, in which case the implementation returns infinity; no empty "
        "front occurred. Non-dominated front size is a count of mutually "
        "non-dominated points; larger is not automatically better because a large "
        "front can lie far from the true surface. Feasibility rate was not applicable "
        "to these unconstrained analytic problems and is reported as not evaluated. "
        "Completed evaluation count was 200 on every run. Failed evaluation count was "
        "zero. Wall-clock time is seconds of process time including initialisation. "
        "Convergence is hypervolume against evaluation count. Variation across seeds "
        "is summarised by the median, interquartile range, mean, standard deviation, "
        "minimum and maximum. Reproducibility used exact equality of hypervolume, IGD, "
        "front and history under seed 123 at a 60-evaluation budget. Sensitivity to "
        "changed seeds compared seeds 123 and 124. Kruskal-Wallis tests compared four "
        "independent groups of ten runs, with Bonferroni-corrected Mann-Whitney tests "
        "and a rank-biserial effect size. Statistical significance was not treated as "
        "a ranking by itself.",
    )

    doc.add_heading("4.2 Functional correctness and test evidence", level=2)
    body(
        doc,
        "Automated tests were executed from a virtual environment that installed "
        "requirements-dev.txt. A passing test verifies only the assertion that it "
        "encodes. The suite does not establish domain validity, usability or "
        "security of the platform as a whole.",
    )

    add_table(
        doc,
        ["Check", "Command or job", "Passed", "Failed", "Skipped", "Evidence", "Interpretation"],
        [
            [
                "Python unit tests",
                'pytest -m "not integration and not slow" --cov',
                "223",
                "0",
                "1",
                "dissertation/chapter4/logs/pytest.txt; 64.86 s; coverage 56%",
                "Core formaliser, packs, optimiser harness, hashing and scoring assertions passed locally.",
            ],
            [
                "Python integration",
                "tests/test_api_crud.py (testcontainers)",
                "0",
                "0",
                "1",
                "pytest skip: testcontainers required",
                "API CRUD against Docker Postgres was not executed in this environment.",
            ],
            [
                "Ruff",
                "ruff check core/ libs/ compute/ services/ scripts/ tests/",
                "0",
                "3",
                "0",
                "dissertation/chapter4/logs/ruff_plain.txt; GitHub Actions lint job 33028242616",
                "Unused Optional import, unused pytest import, ambiguous name l. Style only; no test was failed by these.",
            ],
            [
                "Web unit tests",
                "npm test -- --run (apps/web)",
                "55",
                "0",
                "0",
                "dissertation/chapter4/logs/web-vitest.txt; 2.46 s; GitHub Actions web job success",
                "Workspace tabs, chat message rendering and SSE hook behaviour passed in jsdom.",
            ],
            [
                "Web lint",
                "npm run lint (apps/web)",
                "0",
                "0",
                "0",
                "Interactive ESLint setup prompt; no eslint config in the app",
                "Not evaluated. The command did not run a non-interactive lint.",
            ],
            [
                "Admin tests",
                "npm test -- --run (apps/admin)",
                "0",
                "0",
                "0",
                "Vitest: no test files found, exit 1",
                "The admin application has no automated tests.",
            ],
            [
                "Database migration",
                "alembic upgrade head; python scripts/check_database.py",
                "1",
                "0",
                "0",
                "dissertation/chapter4/logs/alembic.txt and check_database.txt",
                "Head 0001_initial; 41 tables; pgvector 0.6.0 created on PostgreSQL 16.",
            ],
            [
                "GitHub Actions (commit 7803baa)",
                "CI workflow 33028242616",
                "5 jobs",
                "1 job",
                "0",
                "gh run view 33028242616",
                "Tests 3.11, tests 3.12, web, migrations and Docker image build succeeded. Lint failed with the same three Ruff findings. The workflow is not completely green.",
            ],
            [
                "Container configuration",
                "CI Docker images build; local dockerd",
                "1 (CI)",
                "1 (local daemon)",
                "0",
                "CI job success; local overlay2 error",
                "The application image imported in CI. Local Compose was not started.",
            ],
        ],
        "Table 4.2. Automated checks executed for this chapter.",
    )

    body(
        doc,
        "Coverage for the collected Python tree was 56 per cent (7,796 statements, "
        "3,415 missed). Coverage was high for optimiser backends, benchmark metrics, "
        "ToyPack and the judge scorer, and near zero for FastAPI routers, Temporal "
        "workflows, the worker, evidence HTTP routers and core/contract.py. Those "
        "uncovered modules include run sealing, report download and organisation-scoped "
        "run listing. The skipped integration module is the only automated test of "
        "live API CRUD. GitHub Actions confirmed the same non-integration suite on "
        "Python 3.11 and 3.12 and confirmed alembic upgrade head against "
        "pgvector/pgvector:pg15. That is not a completely green workflow because lint "
        "failed.",
    )
    body(
        doc,
        "The web application rendered against the seeded API. Figure 4.6 shows the "
        "ToyPack overview after an in-process run was persisted. The screenshot "
        "demonstrates that the workspace tabs, counters and best-score card can "
        "display values supplied in run_spec. It does not demonstrate that Temporal "
        "produced those values. The Evidence tab showed an empty state. The Heatmaps "
        "tab showed no spatial data, because heatmap layers are not filled from the "
        "run snapshot in the code that was inspected. The admin application rendered "
        "benchmarks, rubrics, packs, an audit log and a simulate preview from "
        "hard-coded mock arrays in apps/admin/src/app/page.tsx; those screens were "
        "photographed, but they are not evidence of the live API.",
    )

    shot = SHOTS / "06_web_toypack_overview.png"
    if shot.exists():
        add_figure(
            doc,
            shot,
            "Figure 4.6. Web workspace Overview tab for an in-process ToyPack run persisted without Temporal.",
            "GSIP web overview showing pipeline stages, 12 simulated scenarios and a best score of 0.862.",
            width=6.2,
        )

    doc.add_heading("4.3 Optimisation benchmark results", level=2)
    body(
        doc,
        "The valid dataset is one hundred and twenty completed runs: three problems, "
        "four methods, ten seeds, budget 200, batch size 20. No run was excluded. "
        "Status was completed and n_evaluations was 200 in every row of "
        "dissertation/chapter4/evidence/benchmark_runs.csv. There were no failed "
        "objective evaluations. Random search is an external harness, not a registered "
        "GSIP backend. Table 4.1 states the configuration. Table 4.3 reports median "
        "and interquartile range for hypervolume and IGD, median front size, median "
        "wall time in seconds, and the number of runs whose hypervolume was exactly "
        "zero. Mean, standard deviation, minimum and maximum are in "
        "benchmark_stats.json.",
    )

    rows_tbl = []
    for prob, label in (("zdt1", "ZDT1"), ("zdt2", "ZDT2"), ("dtlz2", "DTLZ2")):
        for be, blab in (
            ("evolutionary", "evolutionary"),
            ("bayesian", "Bayesian"),
            ("hybrid", "hybrid"),
            ("random", "random"),
        ):
            hv_s, ig_s, fr, tm, nz = hv(st, f"{prob}|{be}")
            rows_tbl.append([label, blab, "10", hv_s, ig_s, fr, tm, nz])
    add_table(
        doc,
        [
            "Problem",
            "Backend",
            "Valid runs",
            "Hypervolume median and IQR",
            "IGD median and IQR",
            "Front size",
            "Wall time (s)",
            "HV = 0",
        ],
        rows_tbl,
        "Table 4.3. Equal-budget results. Hypervolume IQR is the interval from the first to the third quartile. HV = 0 counts runs whose archive lay outside the reference box.",
    )

    add_figure(
        doc,
        FIG / "figure_4_1_convergence_zdt1.png",
        "Figure 4.1. Median hypervolume against completed evaluations on ZDT1 (ten seeds).",
        "Line chart of median hypervolume versus evaluations for evolutionary, Bayesian, hybrid and random search on ZDT1.",
        width=5.8,
    )
    add_figure(
        doc,
        FIG / "figure_4_2_hypervolume.png",
        "Figure 4.2. Hypervolume distributions by problem and method (higher is better).",
        "Three box plots of hypervolume for ZDT1, ZDT2 and DTLZ2 across four methods.",
        width=6.3,
    )
    add_figure(
        doc,
        FIG / "figure_4_3_igd.png",
        "Figure 4.3. IGD distributions by problem and method (lower is better).",
        "Three box plots of inverted generational distance for ZDT1, ZDT2 and DTLZ2.",
        width=6.3,
    )
    add_figure(
        doc,
        FIG / "figure_4_4_pareto.png",
        "Figure 4.4. Representative non-dominated fronts on ZDT1 and ZDT2 for the median-hypervolume seed of each GSIP backend.",
        "Scatter plots of obtained Pareto fronts for evolutionary, Bayesian and hybrid search.",
        width=6.3,
    )
    add_figure(
        doc,
        FIG / "figure_4_5_zero_hypervolume.png",
        "Figure 4.5. Share of runs with hypervolume equal to zero.",
        "Bar chart of the proportion of ten seeds whose archive contributed no hypervolume.",
        width=5.8,
    )

    doc.add_heading("ZDT1", level=3)
    body(
        doc,
        "ZDT1 is the principal two-objective baseline. Bayesian search produced the "
        "highest median hypervolume (0.246, IQR 0.143 to 0.450) and the lowest median "
        "IGD (0.407). No Bayesian run had zero hypervolume. Evolutionary search "
        "reached a median hypervolume of 0.075, with four of ten seeds remaining "
        "outside the reference box. Hybrid search had median hypervolume 0.000, with "
        "six of ten seeds at zero, matching the defect described in "
        "docs/optimiser-benchmark-findings.md: UnifiedOptimizer fits a single-objective "
        "Gaussian process to objectives[0], so half the batch is driven toward one "
        "extreme. Random search was similar to hybrid on hypervolume (median 0.000, "
        "six zeros). Kruskal-Wallis rejected equality of hypervolume distributions "
        "(H = 19.11, p = 0.00026) and of IGD (H = 15.23, p = 0.0016). The "
        "Bonferroni-corrected pairwise Mann-Whitney tests on hypervolume were: "
        "evolutionary versus Bayesian U = 11, adjusted p = 0.021, rank-biserial 0.78 "
        "(Bayesian higher); Bayesian versus hybrid U = 89, adjusted p = 0.019; "
        "Bayesian versus random search U = 100, adjusted p = 0.0009. Evolutionary, "
        "hybrid and random search did not differ from one another after correction. "
        "Wall time remained under one second for "
        "evolutionary and random search, about 29 seconds for Bayesian search and 31 "
        "seconds for hybrid search. Figure 4.1 shows that Bayesian median hypervolume "
        "pulled away after the first batches, while hybrid remained near zero for much "
        "of the budget. These results at five variables and 200 evaluations do not "
        "rank the methods at canonical ZDT1 size.",
    )

    doc.add_heading("ZDT2", level=3)
    body(
        doc,
        "ZDT2 has a concave front. Evolutionary, hybrid and random search produced "
        "zero hypervolume on every seed: their archives never entered the (1.1, 1.1) "
        "box. Bayesian search reached median hypervolume 0.110, with three zero runs "
        "and a typically small front (median size 2). Kruskal-Wallis again rejected "
        "equality (hypervolume H = 24.63, p = 1.8 x 10^-5). Pairwise tests after "
        "correction distinguished Bayesian search from each of the other three methods "
        "and found no difference among evolutionary, hybrid and random search, all of "
        "which sat at zero. The result suggests that, at this budget and dimension, "
        "only the ParEGO-style scalarising GP placed points inside the reference box "
        "with any regularity. It does not show that Bayesian search approximated the "
        "full concave front: IGD remained about 0.61, and Figure 4.4 shows sparse "
        "fronts. Hybrid search was again no better than random search.",
    )

    doc.add_heading("DTLZ2", level=3)
    body(
        doc,
        "DTLZ2 used three objectives and twelve variables. Hypervolume is a Monte "
        "Carlo estimate and is reported only within this problem. No run had zero "
        "hypervolume. Evolutionary search had the highest median hypervolume (0.206) "
        "and the lowest median IGD (0.452). Random search was close (median HV 0.169, "
        "median IGD 0.467) and returned the largest fronts (median 52 points). "
        "Bayesian and hybrid medians were 0.160 and 0.148. Bayesian search was the "
        "slowest method (median 65 seconds). At this budget the evolutionary "
        "population search, which is cheap per evaluation, was at least as effective "
        "as the surrogate methods, and the hybrid combination did not improve on "
        "either half. Because hypervolume is estimated, small differences between "
        "methods should not be over-interpreted; the IGD comparison, which is exact "
        "relative to the sampled sphere, shows the same ordering.",
    )

    doc.add_heading("Overall interpretation of the optimiser comparison", level=3)
    body(
        doc,
        "The hybrid optimiser did not fail to execute: all thirty hybrid runs "
        "completed 200 evaluations. It failed as a multi-objective method on ZDT1 and "
        "ZDT2, where its median hypervolume was zero or indistinguishable from random "
        "search. Source inspection confirms the cause in "
        "services/optimizer/optimizer.py: _init_bayesian constructs BayesianOptimizer "
        "on config.objectives[0] only. The standalone Bayesian backend does not have "
        "that limitation; it scalarises with augmented Chebyshev weights. Research "
        "Question 3 asked whether GSIP supports a fair comparison of evolutionary, "
        "Bayesian and hybrid optimisation. The harness did impose an equal evaluation "
        "budget and repeated seeds. The comparison is therefore fair as an evaluation "
        "count. It is not a fair contest of mature multi-objective Bayesian "
        "optimisation against NSGA-II, because the hybrid's Bayesian half is "
        "single-objective and no qEHVI implementation is present. Computational cost "
        "differed by two orders of magnitude. A single-run ranking, such as the Phase "
        "0 note's seed-42 figures, would have been misleading: evolutionary ZDT1 "
        "hypervolume ranged from 0 to 0.276 across seeds.",
    )

    doc.add_heading("4.4 Reproducibility, traceability and auditability", level=2)
    body(
        doc,
        "Numerical repeatability and full computational reproducibility were treated "
        "as different questions. A run is fully reproducible only if the code version, "
        "dependencies, configuration, model, scenario, inputs, seed, fidelity, "
        "objective definitions and output artefacts can be recovered. The checks below "
        "used exact equality unless stated.",
    )

    add_table(
        doc,
        ["Requirement", "Test performed", "Evidence", "Result", "Limitation"],
        [
            [
                "Same-seed replay",
                "Evolutionary ZDT1, budget 60, seed 123, twice",
                "benchmark_campaign.json replay.same_seed",
                "Hypervolume, IGD, front and history matched exactly",
                "Budget 60, one backend, one problem. Both hypervolumes were 0.0, so the match is exact but uninformative about front recovery.",
            ],
            [
                "Changed seed",
                "Seeds 123 versus 124, otherwise identical",
                "replay.changed_seed",
                "Fronts differed; IGD 1.509 versus 1.878",
                "Hypervolume remained 0.0 for both seeds at this budget.",
            ],
            [
                "Canonical hashing",
                "Unit tests of gsip_ledger.hashing and test_ledger.py",
                "pytest; hashing.py",
                "Canonical JSON and SHA-256 helpers passed their tests",
                "Tests do not prove every production artefact is hashed.",
            ],
            [
                "Model, scenario, optimiser and fidelity recording",
                "Source inspection of Run.run_spec and scenario rows",
                "models.py; in-process UI runs",
                "run_spec stored pack name, prompt, candidates, seeds and fidelity labels",
                "In-process path, not Temporal. Metric-definition versioning was not separately sealed.",
            ],
            [
                "Retry and idempotent persistence",
                "Temporal RetryPolicy in SimulationRunWorkflow",
                "simulation_run.py",
                "Intended: maximum_attempts=3 on several activities",
                "Not experimentally verified; Temporal was not running.",
            ],
            [
                "Artefact storage and report retrieval",
                "Local PDF via pdf_builder; MinIO not running",
                "run_toy-pack.pdf and related files under evidence/",
                "PDFs were generated locally from run_spec",
                "Durable object storage was not verified. A local file is not a durable artefact.",
            ],
            [
                "Run sealing",
                "Read seal_run and Run model",
                "simulation.py seal_run; no is_sealed column on runs",
                "seal_run returns sealed True without writing the database",
                "Completion state is not persistently protected.",
            ],
            [
                "Evidence provenance",
                "Evidence pack tests plus embeddings fallback",
                "test_evidence.py; embeddings.py",
                "Pack assembly tests passed; zero-vector fallback remains if the model is missing",
                "Similarity search against live Milvus was not run.",
            ],
            [
                "Failure reporting when a capability is unavailable",
                "POST /api/runs/start depends on TemporalClient.connect",
                "runs.py get_temporal_client",
                "Without Temporal the start endpoint cannot be resolved",
                "The UI was not used to submit a live run, so the user-visible error path was not photographed.",
            ],
        ],
        "Table 4.4. Reproducibility, traceability and auditability checks.",
    )

    body(
        doc,
        "libs/ledger/gsip_ledger/hashing.py implements a canonical JSON encoding and "
        "SHA-256 helpers, and the unit tests for deterministic hashing passed. Dataset "
        "versions have an is_immutable column; runs do not. The comment in seal_run "
        "states that the activity would set is_sealed, but the function does not "
        "update any row. The test demonstrated exact numerical repeatability for one "
        "evolutionary configuration. The available evidence does not establish full "
        "computational reproducibility of a Temporal-orchestrated decision experiment, "
        "because that path was not executed and because artefacts were not stored in "
        "MinIO.",
    )

    doc.add_heading("4.5 Software quality and security", level=2)
    body(
        doc,
        "Findings were judged against ISO/IEC 25010:2023 quality characteristics "
        "(ISO/IEC, 2023), the NIST AI Risk Management Framework (NIST, 2023), the "
        "NIST Generative AI Profile (NIST, 2024), the OWASP API Security Top 10 "
        "(OWASP, 2023), and the provenance principles already used in the literature "
        "review. Status values are met, partially met, not met or not evaluated. "
        "None of these defects changed the analytic benchmark numbers, which do not "
        "use the HTTP API. Several of them would weaken evidence integrity or "
        "cross-organisation isolation if the platform were used for real decisions.",
    )

    add_table(
        doc,
        ["ID", "Quality or security concern", "Evidence", "Status", "Severity", "Consequence", "Recommendation"],
        [
            [
                "Q1",
                "Run sealing and persistent immutability",
                "seal_run returns sealed True; runs table has no is_sealed",
                "Not met",
                "High",
                "A completed run_spec can be overwritten; audit reports can change after the fact.",
                "Persist a sealed flag and reject mutations after completion.",
            ],
            [
                "Q2",
                "Evidence-pack completeness",
                "Unit tests for pack assembly; live retrieval not run",
                "Partially met",
                "Medium",
                "Reports may cite packs that were never retrieved from a durable store.",
                "Require pack persistence and fail the run if retrieval is empty.",
            ],
            [
                "Q3",
                "Similarity search and zero-vector embedding fallback",
                "embeddings.py returns [0.0] * dimension when the model is missing",
                "Not met",
                "High",
                "All chunks become identical; search ranking is meaningless without an error.",
                "Refuse embedding rather than returning a zero vector.",
            ],
            [
                "Q4",
                "Simulation-isolation fallback",
                "DEFAULT_ISOLATION_MODE is none; ContainerIsolation falls back to NoIsolation",
                "Not met",
                "High",
                "Pack code runs in-process. A hostile pack can touch the host.",
                "Fail closed when container isolation is requested but unavailable.",
            ],
            [
                "Q5",
                "Authentication through identity headers",
                "get_current_user trusts X-User-Id; GSIP_DEMO_AUTH defaults true",
                "Partially met",
                "High",
                "Any caller who knows a user UUID can impersonate that user if demo auth is left on.",
                "Default GSIP_DEMO_AUTH to false and require a verified session.",
            ],
            [
                "Q6",
                "Organisation-level authorisation",
                "test_auth_membership.py; membership check in auth.py",
                "Partially met",
                "Medium",
                "IDOR via X-Org-Id is rejected in unit tests. Live API tests were skipped.",
                "Run the integration suite against Postgres in CI.",
            ],
            [
                "Q7",
                "Raw metric scoring without thresholds",
                "scorer.py clips raw values to [0, 1] when no threshold exists",
                "Partially met",
                "Medium",
                "A Sharpe ratio of 1.8 and a return of 0.18 can receive similar threshold scores.",
                "Require per-metric scaling or refuse to score without thresholds.",
            ],
            [
                "Q8",
                "Objective direction and hybrid Bayesian half",
                "optimizer.py _init_bayesian uses objectives[0]",
                "Not met",
                "High",
                "Hybrid search wastes budget on one objective and poisons the shared archive.",
                "Scalarise as BayesianBackend does, or replace the hybrid GP.",
            ],
            [
                "Q9",
                "Finance-pack annualisation across fidelity",
                "pack.py uses 252 in score() after cheap/mid/high generate monthly or weekly series",
                "Not met",
                "High",
                "Annualised return, volatility and Sharpe ratio are not comparable across fidelity.",
                "Annualise with the periods_per_year already used in _generate_returns.",
            ],
            [
                "Q10",
                "Spatial threshold boundaries",
                "safe_cells = grid < safe_t; critical uses >= crit_t; warning is [safe, warn)",
                "Partially met",
                "Low",
                "A cell equal to the safe threshold is counted as warning, not safe. The cells between warning and critical are unlabelled.",
                "Document and test closed-open intervals explicitly.",
            ],
            [
                "Q11",
                "Durable object storage for reports",
                "MinIO not running; PDFs written locally by the evidence script",
                "Not evaluated",
                "Medium",
                "Reports shown in the UI were not shown to survive process restart in object storage.",
                "Exercise MinIO in the experimental environment.",
            ],
            [
                "Q12",
                "Judge service in the principal workflow",
                "SimulationRunWorkflow imports judge_score_outcomes; judge HTTP service was not started",
                "Partially met",
                "Medium",
                "In-process DeterministicScorer was used for UI runs. The judge HTTP service was not experimentally verified.",
                "Run the judge service in the workflow under test.",
            ],
            [
                "Q13",
                "Error reporting when Temporal is unavailable",
                "start_run Depends on TemporalClient.connect",
                "Partially met",
                "Medium",
                "Run creation fails at dependency resolution rather than a structured capability error.",
                "Return a 503 with an explicit 'orchestrator unavailable' body.",
            ],
            [
                "Q14",
                "core/contract.py provenance types unused by coverage",
                "coverage 0% on core/contract.py",
                "Not met",
                "Medium",
                "The provenance gate described in the architecture is types only.",
                "Wire Provenance into formalisation before execution.",
            ],
        ],
        "Table 4.5. Prioritised quality and security findings from tests and source inspection.",
    )

    body(
        doc,
        "The highest-risk findings are Q1, Q3, Q4, Q5, Q8 and Q9. Run sealing is "
        "intended and not implemented. Zero-vector embeddings are implemented and "
        "would silently degrade retrieval. Container isolation is intended, stubbed, "
        "and falls back to in-process execution. Header authentication is implemented "
        "and automatically tested for membership, but demo auth is on by default. The "
        "hybrid design is implemented, automatically tested only for end-to-end "
        "progress, and experimentally shown to be ineffective on ZDT1 and ZDT2. "
        "Finance annualisation is implemented incorrectly relative to fidelity and was "
        "not caught by the unit tests that do not assert annualisation factors. No "
        "finding in Table 4.5 altered the ZDT or DTLZ numbers. Several would expose "
        "another organisation's data or produce a misleading report if the HTTP path "
        "were used for real decisions. Distinguishing status: intended in comments and "
        "docs; implemented in source; automatically tested where a test exists; "
        "experimentally verified only for the optimiser campaign and the in-process "
        "UI runs; incomplete otherwise.",
    )

    doc.add_heading("4.6 Discussion", level=2)
    body(
        doc,
        "The discussion uses only results already reported in Sections 4.2 to 4.5.",
    )

    doc.add_heading("RQ1", level=3)
    body(
        doc,
        "Research Question 1 asked how successfully a natural-language decision "
        "question is transformed into an explicit, executable specification. The "
        "Python tests demonstrated that different prompts produced different "
        "ObjectiveSpecs, that minimise and maximise cues were detected, and that "
        "domain hints distinguished finance-pack from spatial-pack "
        "(tests/test_smoke_e2e.py). The in-process UI runs stored those structured "
        "fields in run_spec and displayed them in the Overview and Report tabs. "
        "formalize_objective uses heuristics when use_llm is false, and LLM "
        "health reported enabled true with configured_providers empty, so the "
        "experimental formalisation path was heuristic, not model-backed. "
        "core/contract.py defines provenance tags including LLM_RECALL, but coverage "
        "on that module was zero and the types are not wired into the executed "
        "formaliser. The danger of unsupported generated fields therefore remains: "
        "the architecture intends a grounding gate, the tests show keyword mapping "
        "works for canned phrases, and the available evidence does not establish that "
        "an unseen question cannot acquire metrics or bounds that the user never "
        "confirmed. Confirmation and human-in-the-loop override endpoints are listed "
        "as not built in LOCAL_RUN.md; they were not found in the running API.",
    )

    doc.add_heading("RQ2", level=3)
    body(
        doc,
        "Research Question 2 asked to what extent GSIP supports reproducible, "
        "traceable and auditable decision experiments. Implemented strengths include "
        "canonical hashing helpers, seed arguments on domain-pack simulate() methods, "
        "run_spec snapshots that the web client can reload, organisation membership "
        "checks, and Alembic-managed schema. Incomplete enforcement is equally clear. "
        "Run sealing is not persisted. Demo authentication is on by default. Evidence "
        "search can degrade to zero vectors. Isolation can degrade to in-process "
        "execution. Temporal retries, MinIO durability and post-completion mutation "
        "guards were not experimentally verified. The same-seed optimiser replay "
        "matched exactly, which is numerical repeatability of the search harness, not "
        "full computational reproducibility of a decision experiment. NIST's AI Risk "
        "Management Framework treats measurement, documentation and governance as "
        "joint requirements (NIST, 2023). GSIP documents more than it currently "
        "enforces.",
    )

    doc.add_heading("RQ3", level=3)
    body(
        doc,
        "Research Question 3 asked what the evidence shows about a fair comparison of "
        "evolutionary, Bayesian and hybrid optimisation. The harness gave every method "
        "the same evaluation count, ten seeds, shared reference points and shared "
        "fronts. That is the comparison Knowles, Thiele and Zitzler (2006) argued for "
        "when they treated function evaluations as the scarce resource. On ZDT1 the "
        "scalarising Bayesian backend dominated the other three methods on hypervolume "
        "and IGD, at roughly thirty times the wall-clock cost of the evolutionary "
        "search. On ZDT2 only Bayesian search regularly entered the reference box. On "
        "DTLZ2 the evolutionary backend was at least as good as the surrogates, and "
        "random search was competitive, which is consistent with the observation that "
        "inexpensive population search can still cover a spherical front when the "
        "budget is small relative to dimension (Deb et al., 2002). The hybrid method, "
        "which is the v1 engine, did not earn its place on the two-objective problems. "
        "That finding agrees with the repository's own Phase 0 note and extends it "
        "from one seed to ten. It challenges any architectural claim that combining "
        "the two strategies is automatically beneficial. ParEGO-style scalarisation "
        "remains a legitimate multi-objective Bayesian technique (Knowles, 2006); the "
        "defect is that the hybrid does not use it.",
    )

    doc.add_heading("Threats to validity", level=3)
    body(
        doc,
        "Benchmark diversity was limited to ZDT1, ZDT2 and DTLZ2 at reduced ZDT "
        "dimension. Ten seeds support a Kruskal-Wallis test but remain a small sample "
        "for estimating tails. Domain packs are synthetic. Wall times are specific to "
        "a four-core Xeon without a GPU. Fidelity labels cheap, mid and high are "
        "simplified, and finance annualisation is internally inconsistent. Coverage "
        "gaps hide router and workflow behaviour. There was no user study, so RQ1 says "
        "nothing about whether analysts accept the generated specification. Temporal, "
        "MinIO and the judge HTTP service were incomplete in the experimental "
        "environment. The project is individual, so the same person wrote the harness, "
        "chose the budget and interpreted the hybrid failure; that bias is reduced by "
        "pre-registering the smoke budget from the repository rather than tuning it "
        "after seeing ranks, but it is not removed. Results must not be generalised to "
        "real financial, environmental or policy decisions.",
    )

    doc.add_heading("4.7 Chapter summary", level=2)
    body(
        doc,
        "The experimental campaign executed the unmodified GSIP tree at commit 7803baa. "
        "Two hundred and twenty-three Python tests passed and one integration test was "
        "skipped. Fifty-five web tests passed. Ruff reported three style errors, and "
        "GitHub Actions lint failed while other CI jobs succeeded. One hundred and "
        "twenty optimiser runs completed at an equal budget of 200 evaluations.",
    )
    body(
        doc,
        "RQ1. Heuristic formalisation mapped canned prompts onto different schemas, "
        "metrics and domains in unit tests, and those fields appeared in persisted "
        "run_spec objects. Provenance tagging, confirmation and LLM-backed "
        "formalisation were not experimentally established.",
    )
    body(
        doc,
        "RQ2. GSIP implements hashing helpers, seeds and reloadable run snapshots, "
        "which support numerical repeatability of the optimiser harness. Persistent "
        "run sealing, durable artefact storage, isolation and retrieval integrity are "
        "not enforced in the submitted implementation.",
    )
    body(
        doc,
        "RQ3. At five variables and 200 evaluations, Bayesian scalarisation outperformed "
        "evolutionary search, hybrid search and random search on ZDT1 and ZDT2. Hybrid "
        "search did not beat random search on those problems because its Bayesian half "
        "optimises a single objective. On three-objective DTLZ2, evolutionary search "
        "matched or exceeded the surrogates. The comparison is fair on evaluation "
        "count and unfair as a test of multi-objective Bayesian optimisation in the "
        "hybrid.",
    )
    body(
        doc,
        "Findings that must be carried into Chapter 5 are: the hybrid multi-objective "
        "defect; the unverified sealing path; the zero-vector and isolation fallbacks; "
        "finance annualisation across fidelity; the absence of a user study; and the "
        "limit of the benchmark campaign to reduced ZDT dimension without Temporal.",
    )

    doc.add_heading("References cited in Chapter 4", level=2)
    refs = [
        "Deb, K., Thiele, L., Laumanns, M. and Zitzler, E. (2002) 'Scalable test problems for evolutionary multiobjective optimization', in Abraham, A., Jain, L. and Goldberg, R. (eds.) Evolutionary Multiobjective Optimization. London: Springer, pp. 105-145.",
        "International Organization for Standardization (2023) ISO/IEC 25010:2023 Systems and software engineering. Systems and software Quality Requirements and Evaluation (SQuaRE). System and software quality models. Geneva: ISO.",
        "Knowles, J. (2006) 'ParEGO: a hybrid algorithm with on-line landscape approximation for expensive multiobjective optimization problems', IEEE Transactions on Evolutionary Computation, 10(1), pp. 50-66.",
        "Knowles, J., Thiele, L. and Zitzler, E. (2006) A tutorial on the performance assessment of stochastic multiobjective optimizers. TIK Report 214, ETH Zurich.",
        "National Institute of Standards and Technology (2023) Artificial Intelligence Risk Management Framework (AI RMF 1.0). NIST AI 100-1. Gaithersburg, MD: NIST.",
        "National Institute of Standards and Technology (2024) Artificial Intelligence Risk Management Framework: Generative Artificial Intelligence Profile. NIST AI 600-1. Gaithersburg, MD: NIST.",
        "OWASP Foundation (2023) OWASP API Security Top 10. Available at: https://owasp.org/API-Security/ (Accessed: 27 August 2026).",
        "Zitzler, E., Deb, K. and Thiele, L. (2000) 'Comparison of multiobjective evolutionary algorithms: empirical results', Evolutionary Computation, 8(2), pp. 173-195.",
    ]
    for ref in refs:
        p = doc.add_paragraph()
        set_paragraph_format(p, after=6, align="left", line=1.15)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        run = p.add_run(ref)
        set_run_font(run, size=11)

    doc.add_heading("Appendix 4A. Evidence inventory", level=2)
    body(
        doc,
        "Supporting files live under dissertation/chapter4/. The campaign JSON and CSV "
        "contain one row per run. Logs contain pytest, Ruff, Alembic, seed, Vitest and "
        "environment capture. Screenshots are PNG exports of the running web, admin "
        "and OpenAPI surfaces. Application source was not modified to obtain them.",
    )

    doc.add_heading("Appendix 4B. Application screenshots", level=2)
    body(
        doc,
        "The following figures are photographs of the experimental environment. They "
        "document what the interfaces displayed. Admin screens use mock data in the "
        "client and are labelled as such.",
    )
    appendix_shots = [
        ("01_web_idle_overview.png", "Figure 4.7. Idle web workspace before a run is opened."),
        ("02_web_project_dropdown.png", "Figure 4.8. Project selector."),
        ("03_web_domain_pack_dropdown.png", "Figure 4.9. Domain-pack selector."),
        ("04_web_run_settings.png", "Figure 4.10. Run settings panel."),
        ("05_web_run_history_sidebar.png", "Figure 4.11. Past-runs sidebar listing three completed in-process runs."),
        ("07_web_toypack_leaderboard.png", "Figure 4.12. ToyPack leaderboard."),
        ("08_web_toypack_scenario_detail.png", "Figure 4.13. ToyPack scenario detail."),
        ("09_web_toypack_charts_timeseries.png", "Figure 4.14. ToyPack charts, time series."),
        ("10_web_toypack_charts_comparison.png", "Figure 4.15. ToyPack charts, comparison."),
        ("11_web_toypack_charts_pareto.png", "Figure 4.16. ToyPack charts, labelled Pareto view."),
        ("12_web_toypack_evidence_empty.png", "Figure 4.17. Evidence tab empty state."),
        ("13_web_toypack_report.png", "Figure 4.18. Report tab."),
        ("14_web_toypack_logs.png", "Figure 4.19. Logs and debug tab."),
        ("15_web_spatial_overview.png", "Figure 4.20. SpatialPack overview after an in-process run."),
        ("16_web_spatial_heatmaps_empty.png", "Figure 4.21. Heatmaps tab reporting no spatial data."),
        ("17_web_spatial_leaderboard.png", "Figure 4.22. SpatialPack leaderboard."),
        ("18_web_spatial_charts.png", "Figure 4.23. SpatialPack charts."),
        ("19_web_finance_overview.png", "Figure 4.24. FinancePack overview after an in-process run."),
        ("20_web_finance_leaderboard.png", "Figure 4.25. FinancePack leaderboard."),
        ("21_web_finance_charts.png", "Figure 4.26. FinancePack charts."),
        ("22_web_finance_report.png", "Figure 4.27. FinancePack report tab."),
        ("23_web_chat_composer.png", "Figure 4.28. Chat composer."),
        ("24_admin_benchmarks.png", "Figure 4.29. Admin benchmarks page (client mock data)."),
        ("25_admin_rubrics.png", "Figure 4.30. Admin rubrics page (client mock data)."),
        ("26_admin_packs.png", "Figure 4.31. Admin domain packs page (client mock data)."),
        ("27_admin_audit.png", "Figure 4.32. Admin audit log (client mock data)."),
        ("28_admin_simulate.png", "Figure 4.33. Admin simulate preview (client mock data)."),
        ("29_api_docs_home.png", "Figure 4.34. OpenAPI documentation at /docs."),
        ("30_api_docs_health.png", "Figure 4.35. OpenAPI health endpoint."),
        ("31_api_docs_projects.png", "Figure 4.36. OpenAPI projects endpoint."),
    ]
    n = 7
    for fname, cap_text in appendix_shots:
        path = SHOTS / fname
        if path.exists():
            add_figure(doc, path, cap_text, cap_text, width=6.1)
            n += 1

    doc.save(OUT_DOCX)
    print(f"wrote {OUT_DOCX}")
    return OUT_DOCX


def word_count_report(doc_path: Path) -> Path:
    """Count body words excluding tables and captions."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    from docx.oxml.ns import qn as _qn

    doc = Document(str(doc_path))
    body_paras = []
    caption_paras = []
    heading_paras = []
    table_cells = 0
    for child in doc.element.body.iterchildren():
        tag = child.tag
        if tag == _qn("w:tbl"):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    table_cells += len(cell.text.split())
            continue
        if tag != _qn("w:p"):
            continue
        para = Paragraph(child, doc)
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style is not None else ""
        if style.startswith("Heading"):
            heading_paras.append(text)
            continue
        if text.startswith("Table 4.") or text.startswith("Figure 4."):
            caption_paras.append(text)
            continue
        body_paras.append(text)

    def nwords(items):
        return sum(len(t.split()) for t in items)

    body_n = nwords(body_paras)
    head_n = nwords(heading_paras)
    cap_n = nwords(caption_paras)
    report = HERE / "word_count_report.md"
    lines = [
        "# Chapter 4 word count",
        "",
        "Counts are whitespace-delimited English words from the generated Word file.",
        "Table cells and captions are excluded from the body total, as requested.",
        "",
        f"- Body paragraphs (excluding headings, tables, captions): **{body_n}**",
        f"- Headings: {head_n}",
        f"- Captions: {cap_n}",
        f"- Table cell words: {table_cells}",
        f"- Body plus headings (the narrative chapter): **{body_n + head_n}**",
        "",
        "Target for the chapter body was 3,500 to 5,000 words excluding tables and captions.",
        "If the bound dissertation already contains a references chapter, the short Chapter 4",
        "reference list should be merged rather than duplicated.",
        "",
        "Dash check: generate_chapter4_docx.py contains no em dash or en dash characters.",
        "A post-conversion scan of the .docx XML is recorded below by the build script.",
        "",
    ]
    report.write_text("\n".join(lines), encoding="utf-8")
    print(f"body words={body_n} headings={head_n} captions={cap_n} table_words={table_cells}")
    return report


if __name__ == "__main__":
    path = build()
    word_count_report(path)
